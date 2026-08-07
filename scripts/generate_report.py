# -*- coding: utf-8 -*-
"""
Tạo file Word báo cáo tiến độ tuần 1-7 cho Đồ án Tốt Nghiệp.
Nội dung được viết lại dựa trên project thực tế (mã nguồn trong e:\datn).
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "BaoCaoTienDoHoanThanhDATN-22050040.docx")


def set_cell_shading(cell, color):
    """Tô màu nền cho ô."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_paragraph(doc, text, bold=False, size=12, align=None, space_after=6, font_name="Times New Roman"):
    """Thêm đoạn văn với định dạng."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading_custom(doc, text, size=14, color="000000"):
    """Thêm heading với định dạng tùy chỉnh."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_after = Pt(6)
    return p


def main():
    doc = Document()

    # Thiết lập margin
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # ─── Header ────────────────────────────────────────────────────────
    add_paragraph(doc, "TRƯỜNG ĐẠI HỌC BÌNH DƯƠNG", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_paragraph(doc, "VIỆN TRÍ TUỆ NHÂN TẠO VÀ CHUYỂN ĐỔI SỐ", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_paragraph(doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_paragraph(doc, "Độc lập – Tự do – Hạnh phúc", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_paragraph(doc, "ĐÁNH GIÁ TIẾN ĐỘ GIỮA GIAI ĐOẠN ĐỒ ÁN TỐT NGHIỆP", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # ─── 1. Thông tin đề tài ───────────────────────────────────────────
    add_heading_custom(doc, "1. Thông tin đề tài:", size=14)

    add_paragraph(doc, "TÊN ĐỀ TÀI: PHÁT TRIỂN NỀN TẢNG SMART SHOPPING ASSISTANT TÍCH HỢP PHÂN TÍCH CẢM XÚC VÀ DỰ BÁO XU HƯỚNG GIÁ SẢN PHẨM", bold=True, size=12, space_after=6)
    add_paragraph(doc, "Cán bộ hướng dẫn (CBHD): ThS. Dương Anh Tuấn", size=12, space_after=6)
    add_paragraph(doc, "Thời gian thực hiện: Từ ngày 22/06/2026. Đến ngày 13/09/2026", size=12, space_after=6)
    add_paragraph(doc, "Sinh viên thực hiện:   Nguyễn Hoàng An – 22050040", size=12, space_after=12)

    # ─── 2. Tiến độ thực hiện ──────────────────────────────────────────
    add_heading_custom(doc, "2. Tiến độ thực hiện:", size=14)

    # Tạo bảng tiến độ
    table = doc.add_table(rows=13, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header bảng
    headers = ["Thời gian", "Công việc sinh viên đã đạt", "Nhận xét tiến độ\n(So với đề cương chi tiết)"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "D9E2F3")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"

    # Dữ liệu các tuần - viết lại theo project thực tế
    data = [
        # Tuần 1
        ("Tuần 1\n(ngày 22/06 tới ngày 28/06)",
         "Tìm hiểu về cấu trúc trang, cơ chế render của FPT Shop, Thế Giới Di Động, CellphoneS.\n"
         "Tìm hiểu về RAG pipeline (Retriever + Generator) và cách triển khai vào sản phẩm.\n"
         "Tìm hiểu về mô hình PhoBERT (vinai/phobert-base-v2) cho phân tích cảm xúc tiếng Việt.\n"
         "Tìm hiểu về mạng LSTM cho dự báo chuỗi thời gian giá sản phẩm.\n"
         "Bắt đầu crawl dữ liệu 3 sàn bằng API.",
         "Đúng tiến độ"),
        # Tuần 2
        ("Tuần 2\n(ngày 29/06 tới ngày 5/07)",
         "Crawl được dữ liệu của CellphoneS bằng API và lưu vào MongoDB Compass (collection 'cellphones').\n"
         "Crawl dữ liệu FPT Shop và Thế Giới Di Động bằng Playwright thay thế do API bị chặn.\n"
         "Xây dựng cấu trúc dữ liệu sản phẩm với price_history (lịch sử giá) và comments (bình luận).",
         "Đúng tiến độ"),
        # Tuần 3
        ("Tuần 3\n(ngày 06/07 tới ngày 12/07)",
         "Xây dựng hệ thống Core Backend API bằng FastAPI (main.py) với các endpoint tìm kiếm, lịch sử giá, dự báo.\n"
         "Tiền xử lý dữ liệu và huấn luyện (Fine-tune) mô hình PhoBERT cho phân loại cảm xúc.\n"
         "Xây dựng mô hình LSTM (train_lstm.py) với cơ chế lưu/khôi phục model vào saved_models/lstm.\n"
         "Tiếp tục chỉnh sửa crawl dữ liệu cho FPT Shop và Thế Giới Di Động.",
         "Đúng tiến độ"),
        # Tuần 4
        ("Tuần 4\n(ngày 13/07 tới ngày 19/07)",
         "Hoàn thiện hệ thống xác thực người dùng cho các dịch vụ API (JWT).\n"
         "Xây dựng endpoint quản lý sản phẩm yêu thích và cơ chế lưu trữ lịch sử giá (price_history append-only).\n"
         "Tối ưu hóa các tiến trình nền (BackgroundScheduler) để tránh nghẽn mạch khi cào dữ liệu định kỳ mỗi giờ.\n"
         "Xây dựng cơ chế cache in-memory (TTL 1 giờ) cho kết quả tìm kiếm.",
         "Đúng tiến độ"),
        # Tuần 5
        ("Tuần 5\n(ngày 20/07 tới ngày 26/07)",
         "Thực hiện tiền xử lý tập dữ liệu văn bản bình luận: chuẩn hóa từ lóng, teencode và áp dụng công cụ VnCoreNLP để tách từ tiếng Việt.\n"
         "Tiến hành huấn luyện (Fine-tune) mô hình PhoBERT nhằm phân loại cảm xúc bình luận thành 3 nhãn (Tích cực, Tiêu cực, Trung lập) để quy đổi thành điểm số chất lượng thực tế (RQS).\n"
         "Tối ưu hóa cấu trúc mạng LSTM với các lớp ẩn (64, 32) và Dropout (0.2) để huấn luyện dự báo chuỗi thời gian cho lịch sử giá sản phẩm.",
         "Đang thực hiện"),
        # Tuần 6
        ("Tuần 6\n(ngày 27/07 tới ngày 02/08)",
         "Hoàn thiện mô hình PhoBERT phân loại cảm xúc 3 nhãn và tích hợp vào hệ thống (ai/sentiment).\n"
         "Xây dựng endpoint phân tích cảm xúc bình luận sản phẩm (Sentiment Analysis API) lưu kết quả vào collection 'sentiments'.\n"
         "Tích hợp điểm số chất lượng thực tế (RQS) vào hệ thống khuyến nghị mua hàng (PQS - Product Quality Score).\n"
         "Xây dựng RAG pipeline (ai/rag) kết hợp Retriever truy xuất dữ liệu MongoDB và Generator tạo câu trả lời cho chatbot.",
         "Đúng tiến độ"),
        # Tuần 7
        ("Tuần 7\n(ngày 03/08 tới ngày 09/08)",
         "Hoàn thiện mô hình LSTM dự báo giá với các lớp Dropout và lưu model vào saved_models/lstm theo từng sản phẩm.\n"
         "Xây dựng endpoint dự báo giá sản phẩm (Price Forecast API) và lưu kết quả vào collection 'forecasts'.\n"
         "Xây dựng hệ thống đánh giá độ chính xác dự báo (MAE, RMSE, MAPE, Direction Accuracy) trong evaluate.py.\n"
         "Xây dựng endpoint khuyến nghị mua hàng (Buy Recommendation) kết hợp PQS, thống kê giá và dự báo LSTM.",
         "Đúng tiến độ"),
        # Tuần 8-12 (trống)
        ("Tuần 8\n(ngày 10/08 tới ngày 16/08)", "", ""),
        ("Tuần 9\n(ngày 17/08 tới ngày 23/08)", "", ""),
        ("Tuần 10\n(ngày 24/08 tới ngày 30/08)", "", ""),
        ("Tuần 11\n(ngày 31/08 tới ngày 06/09)", "", ""),
        ("Tuần 12\n(ngày 07/09 tới ngày 13/09)", "", ""),
    ]

    for row_idx, (time, work, note) in enumerate(data, start=1):
        cells = table.rows[row_idx].cells
        cells[0].text = time
        cells[1].text = work
        cells[2].text = note

        # Định dạng font cho từng ô
        for cell in cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        # Căn giữa cột thời gian và nhận xét
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ─── Kết luận ──────────────────────────────────────────────────────
    add_paragraph(doc, "", size=12, space_after=6)
    add_heading_custom(doc, "KẾT LUẬN: đạt/ không đạt tiến độ thực hiện đồ án (Tỷ lệ %).", size=12)

    add_paragraph(doc, "", size=12, space_after=12)
    add_paragraph(doc, "Thành phố Hồ Chí Minh, ngày      tháng 7 năm 2026", size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=24)

    add_paragraph(doc, "GIẢNG VIÊN HƯỚNG DẪN", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)

    # Lưu file
    doc.save(OUTPUT_PATH)
    print("File created:", os.path.abspath(OUTPUT_PATH))


if __name__ == "__main__":
    main()